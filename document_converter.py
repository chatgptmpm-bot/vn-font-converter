"""
Document converter for Office files (DOCX, XLSX)
Converts font encoding while preserving formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border
import os
from vni_unicode_mapping import (
    vni_to_unicode, unicode_to_vni, is_vni_font, is_unicode_font, VNI_FONTS, UNICODE_FONTS
)


class DocumentConverter:
    def __init__(self):
        self.errors = []
    
    def convert_word_document(self, input_path, output_path, conversion_type='auto'):
        """
        Convert Word document (.docx)
        conversion_type: 'auto', 'vni_to_unicode', 'unicode_to_vni'
        """
        try:
            doc = Document(input_path)
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                self._convert_paragraph(paragraph, conversion_type)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._convert_paragraph(paragraph, conversion_type)
            
            # Process headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    self._convert_paragraph(paragraph, conversion_type)
                for paragraph in section.footer.paragraphs:
                    self._convert_paragraph(paragraph, conversion_type)
            
            doc.save(output_path)
            return True, f"Successfully converted to {output_path}"
        
        except Exception as e:
            error_msg = f"Error converting Word document: {str(e)}"
            self.errors.append(error_msg)
            return False, error_msg
    
    def _convert_paragraph(self, paragraph, conversion_type='auto'):
        """Convert text in a paragraph while preserving formatting"""
        for run in paragraph.runs:
            # Detect conversion type if auto
            if conversion_type == 'auto':
                font_name = run.font.name
                if is_vni_font(font_name):
                    conversion_type = 'vni_to_unicode'
                    run.font.name = 'Times New Roman'  # Change to Unicode font
                elif is_unicode_font(font_name):
                    conversion_type = 'unicode_to_vni'
                else:
                    # Try to detect from text content
                    conversion_type = self._detect_encoding(run.text)
            
            # Convert text
            if conversion_type == 'vni_to_unicode':
                run.text = vni_to_unicode(run.text)
                if run.font.name and is_vni_font(run.font.name):
                    run.font.name = 'Times New Roman'
            elif conversion_type == 'unicode_to_vni':
                run.text = unicode_to_vni(run.text)
                if run.font.name:
                    # Change to VNI font
                    run.font.name = 'VNI-Times'
    
    def convert_excel_document(self, input_path, output_path, conversion_type='auto'):
        """
        Convert Excel document (.xlsx)
        conversion_type: 'auto', 'vni_to_unicode', 'unicode_to_vni'
        """
        try:
            wb = load_workbook(input_path)
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            # Detect conversion type if auto
                            current_conversion = conversion_type
                            if current_conversion == 'auto':
                                font_name = cell.font.name if cell.font else None
                                if is_vni_font(font_name):
                                    current_conversion = 'vni_to_unicode'
                                    # Update font
                                    old_font = cell.font
                                    cell.font = Font(
                                        name='Times New Roman',
                                        size=old_font.size if old_font else 11,
                                        bold=old_font.bold if old_font else False,
                                        italic=old_font.italic if old_font else False,
                                        color=old_font.color if old_font else None,
                                    )
                                else:
                                    current_conversion = self._detect_encoding(cell.value)
                            
                            # Convert text
                            if current_conversion == 'vni_to_unicode':
                                cell.value = vni_to_unicode(cell.value)
                            elif current_conversion == 'unicode_to_vni':
                                cell.value = unicode_to_vni(cell.value)
            
            wb.save(output_path)
            return True, f"Successfully converted to {output_path}"
        
        except Exception as e:
            error_msg = f"Error converting Excel document: {str(e)}"
            self.errors.append(error_msg)
            return False, error_msg
    
    def _detect_encoding(self, text):
        """
        Detect if text is VNI or Unicode encoded
        Returns: 'vni_to_unicode', 'unicode_to_vni', or 'none'
        """
        if not text:
            return 'none'
        
        # Check for VNI-specific character (ð used in Vietnamese)
        vni_indicators = ['ð', 'Ð']
        unicode_indicators = ['đ', 'Đ']
        
        vni_count = sum(1 for c in text if c in vni_indicators)
        unicode_count = sum(1 for c in text if c in unicode_indicators)
        
        if vni_count > unicode_count:
            return 'vni_to_unicode'
        elif unicode_count > vni_count:
            return 'unicode_to_vni'
        
        return 'none'
    
    def convert_file(self, input_path, output_path=None, conversion_type='auto'):
        """
        Auto-detect file type and convert
        Returns: (success: bool, message: str, output_path: str)
        """
        if not os.path.exists(input_path):
            return False, f"File not found: {input_path}", None
        
        # Generate output path if not provided
        if not output_path:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_converted{ext}"
        
        # Detect file type
        _, ext = os.path.splitext(input_path)
        ext = ext.lower()
        
        if ext == '.docx':
            success, message = self.convert_word_document(input_path, output_path, conversion_type)
        elif ext == '.xlsx':
            success, message = self.convert_excel_document(input_path, output_path, conversion_type)
        else:
            return False, f"Unsupported file format: {ext}", None
        
        if success:
            return True, message, output_path
        else:
            return False, message, None
